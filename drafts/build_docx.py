#!/usr/bin/env python3
"""Generate one docx per article from markdown files."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ARTICLES_DIR = Path(__file__).parent.parent / "tasks" / "knowledge-map" / "reorg" / "articles"
OUTPUT_DIR = Path(__file__).parent / "docx"


def add_formatted_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph(style=style)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|_.*?_)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)
    return p


def make_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.5
    h0_style = doc.styles['Title']
    h0_style.font.color.rgb = RGBColor(0xD2, 0x9A, 0x38)
    h2_style = doc.styles['Heading 2']
    h2_style.font.size = Pt(14)
    for section in doc.sections:
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    return doc


def process_markdown(doc, md_path):
    text = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')
    in_blockquote = False
    blockquote_lines = []
    in_table_header = True

    def flush_bq():
        nonlocal in_blockquote, blockquote_lines
        if blockquote_lines:
            bq_text = ' '.join(blockquote_lines)
            p = doc.add_paragraph(style='Intense Quote')
            p.add_run(bq_text)
            blockquote_lines = []
        in_blockquote = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith('*Next:'):
            continue

        if stripped.startswith('# ') and not stripped.startswith('## '):
            title = stripped[2:]
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            byline = doc.add_paragraph('By Zzy')
            byline.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            byline.runs[0].font.size = Pt(10)
            draft = doc.add_paragraph('DRAFT')
            draft.runs[0].font.color.rgb = RGBColor(0xD2, 0x9A, 0x38)
            draft.runs[0].font.size = Pt(9)
            doc.add_paragraph('')
            continue

        if stripped.startswith('### '):
            flush_bq()
            doc.add_heading(stripped[4:], level=3)
            continue

        if stripped.startswith('## '):
            flush_bq()
            doc.add_heading(stripped[3:], level=2)
            continue

        if stripped.startswith('>'):
            in_blockquote = True
            blockquote_lines.append(stripped[1:].strip())
            continue
        elif in_blockquote:
            flush_bq()

        if stripped.startswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if all(set(c) <= set('- :') for c in cells):
                continue
            row_text = '  |  '.join(cells)
            add_formatted_paragraph(doc, row_text)
            continue

        if stripped.startswith('- '):
            add_formatted_paragraph(doc, stripped[2:], style='List Bullet')
            continue

        if not stripped:
            continue

        add_formatted_paragraph(doc, stripped)

    flush_bq()


def main():
    OUTPUT_DIR.mkdir(exist_ok=True)

    articles = sorted(ARTICLES_DIR.glob('*.md'))
    for article_path in articles:
        doc = make_doc()
        process_markdown(doc, article_path)
        out_name = article_path.stem + '.docx'
        out_path = OUTPUT_DIR / out_name
        doc.save(str(out_path))
        print(f"  {out_name}")

    print(f"\nGenerated {len(articles)} files in {OUTPUT_DIR}/")


if __name__ == '__main__':
    main()
